"""
File parsing utilities for Nova HR Assistant

This module provides functionality to extract text from various file formats
including PDF, Word documents, and plain text files.
"""

import os
import logging
from pathlib import Path
from typing import Optional

try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


logger = logging.getLogger(__name__)


class FileParsingError(Exception):
    """Exception raised when file parsing fails"""
    pass


class FileParser:
    """
    Utility class for parsing various file formats to extract text content
    """
    
    SUPPORTED_EXTENSIONS = {'.txt', '.pdf', '.docx', '.doc'}
    
    @classmethod
    def is_supported_file(cls, file_path: str) -> bool:
        """
        Check if the file type is supported
        
        Args:
            file_path: Path to the file
            
        Returns:
            True if file type is supported, False otherwise
        """
        extension = Path(file_path).suffix.lower()
        return extension in cls.SUPPORTED_EXTENSIONS
    
    @classmethod
    def extract_text(cls, file_path: str) -> str:
        """
        Extract text content from a file
        
        Args:
            file_path: Path to the file to parse
            
        Returns:
            Extracted text content
            
        Raises:
            FileParsingError: If file parsing fails
            FileNotFoundError: If file doesn't exist
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        logger.info(f"Parsing file: {file_path} (type: {extension})")
        
        try:
            if extension == '.txt':
                return cls._parse_text_file(file_path)
            elif extension == '.pdf':
                return cls._parse_pdf_file(file_path)
            elif extension in ['.docx', '.doc']:
                return cls._parse_word_file(file_path)
            else:
                raise FileParsingError(f"Unsupported file type: {extension}")
                
        except Exception as e:
            logger.error(f"Failed to parse file {file_path}: {str(e)}")
            raise FileParsingError(f"Failed to parse {file_path}: {str(e)}")
    
    @staticmethod
    def _parse_text_file(file_path: Path) -> str:
        """
        Parse plain text file
        
        Args:
            file_path: Path to text file
            
        Returns:
            File content as string
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as f:
                content = f.read()
        
        if not content.strip():
            raise FileParsingError("Text file is empty")
        
        return content.strip()
    
    @staticmethod
    def _parse_pdf_file(file_path: Path) -> str:
        """
        Parse PDF file using multiple methods for better compatibility
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text content
        """
        if not PDF_AVAILABLE:
            raise FileParsingError("PDF parsing libraries not installed. Install with: pip install PyPDF2 pdfplumber")
        
        text_content = ""
        
        # Try pdfplumber first (better for complex layouts)
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n"
            
            if text_content.strip():
                logger.info(f"Successfully extracted text using pdfplumber: {len(text_content)} characters")
                return text_content.strip()
        except Exception as e:
            logger.warning(f"pdfplumber failed for {file_path}: {str(e)}")
        
        # Fallback to PyPDF2
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n"
            
            if text_content.strip():
                logger.info(f"Successfully extracted text using PyPDF2: {len(text_content)} characters")
                return text_content.strip()
        except Exception as e:
            logger.warning(f"PyPDF2 failed for {file_path}: {str(e)}")
        
        if not text_content.strip():
            raise FileParsingError("Could not extract text from PDF file. The PDF might be image-based or corrupted.")
        
        return text_content.strip()
    
    @staticmethod
    def _parse_word_file(file_path: Path) -> str:
        """
        Parse Word document file
        
        Args:
            file_path: Path to Word document
            
        Returns:
            Extracted text content
        """
        if not DOCX_AVAILABLE:
            raise FileParsingError("Word document parsing library not installed. Install with: pip install python-docx")
        
        extension = file_path.suffix.lower()
        
        if extension == '.doc':
            raise FileParsingError("Legacy .doc files are not supported. Please convert to .docx format.")
        
        try:
            doc = Document(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text.strip())
            
            full_text = "\n".join(text_content)
            
            if not full_text.strip():
                raise FileParsingError("Word document appears to be empty")
            
            logger.info(f"Successfully extracted text from Word document: {len(full_text)} characters")
            return full_text.strip()
            
        except Exception as e:
            raise FileParsingError(f"Failed to parse Word document: {str(e)}")
    
    @classmethod
    def get_file_info(cls, file_path: str) -> dict:
        """
        Get information about a file
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dictionary with file information
        """
        if not os.path.exists(file_path):
            return {"exists": False}
        
        file_path = Path(file_path)
        stat = file_path.stat()
        
        return {
            "exists": True,
            "name": file_path.name,
            "extension": file_path.suffix.lower(),
            "size_bytes": stat.st_size,
            "size_mb": round(stat.st_size / (1024 * 1024), 2),
            "is_supported": cls.is_supported_file(str(file_path)),
            "absolute_path": str(file_path.absolute())
        }
    
    @classmethod
    def validate_file(cls, file_path: str) -> list:
        """
        Validate a file for parsing
        
        Args:
            file_path: Path to the file
            
        Returns:
            List of validation errors (empty if valid)
        """
        errors = []
        
        if not os.path.exists(file_path):
            errors.append(f"File does not exist: {file_path}")
            return errors
        
        file_info = cls.get_file_info(file_path)
        
        if not file_info["is_supported"]:
            supported = ", ".join(cls.SUPPORTED_EXTENSIONS)
            errors.append(f"Unsupported file type: {file_info['extension']}. Supported types: {supported}")
        
        if file_info["size_bytes"] == 0:
            errors.append("File is empty")
        
        if file_info["size_mb"] > 50:  # 50MB limit
            errors.append(f"File is too large: {file_info['size_mb']}MB (max 50MB)")
        
        # Check file permissions
        try:
            with open(file_path, 'rb') as f:
                f.read(1)
        except PermissionError:
            errors.append("Permission denied: cannot read file")
        except Exception as e:
            errors.append(f"Cannot access file: {str(e)}")
        
        return errors